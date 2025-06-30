import google.generativeai as genai
import os

# IA
genai.configure(api_key="sua api do gemini aqui")

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from datetime import datetime

from arquivos_base import arquivo_validacao_word

# Arquivos de Base
from arquivos_base import clientes_oferta, produtos_oferta, clientes_com_pendencia

# Arquivos das validações
from ValidacoesPY.Verificacao_1_CNPJ import validar_clientes, qtde_divergente_01
from ValidacoesPY.Verificacao_2_Somente_ProdutosOferta import validar_produtos, qtde_divergente_02
from ValidacoesPY.Verificacao_3_Pendencia_Juridica import validar_clientes_pendencia, qtde_divergente_03


caminho_word_base = r"C:\fakepath\AUTOMACAO_APURACAO\Integracao\word"
arquivo_word_base = "Laudo Apuração BASE.docx"


def formata_titulo(doc, nome_titulo, level):
    """
    Adiciona um título formatado ao documento.
    """
    titulo = doc.add_heading(nome_titulo, level=level)
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = titulo.runs[0]
    run.font.name = "Calibri"
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)


def formata_paragrafo(doc, texto):
    """
    Adiciona um parágrafo formatado e justificado ao documento.
    """
    par = doc.add_paragraph()
    run = par.add_run(texto)

    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 0, 0)

    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# 1°
introducao = {
    "INTRODUÇÃO": "Laudo de apuração de Oferta, que se trata das validações de CNPJ Participantes, Produtos, Pendência Juridica, referente a Oferta-Junho 2025."
}

# 2°
lista_dados_utilizados = [clientes_oferta, produtos_oferta, clientes_com_pendencia]

# 3°
texto_validacoes = {
    "CNPJ PARTICIPANTES": "Verificar se os clientes configurados no banco estão de acordo com o arquivo enviado pelo Cliente.",
    "VERIFICAR PRODUTOS": "Veriicar se os produtos vindo do arquivo enviado pelo cliente estão, estão de acordo com o que está cadastrado no banco de dados.",
    "PENDÊNCIA JURIDICA": "Verificar se algum CNPJ que está dentro das Ofertas estão com Pendência Juridicas.",
}

status_item = {
    "Verificar CNPJ": qtde_divergente_01,
    "Verificar Produtos": qtde_divergente_02,
    "Pendência Júridica": qtde_divergente_03
}

status_text = [
    "Correto",
    "Divergente"
]

validacao = []
status = []
qtde_divergente = []


def gerar_resumo_IA(resumo_validacao):
    lista_resultados = []
    for item, qtde in resumo_validacao.items():
        if qtde > 0:
            lista_resultados.append(f"{item}: {qtde} divergência(s)")
        else:
            lista_resultados.append(f"{item}: sem divergência")

    texto_resultados = "\n".join(lista_resultados)

    prompt = f"""
            Você é um assistente que escreve resumos técnicos de análises de dados para relatórios profissionais.

            Abaixo estão os resultados de uma apuração de validações de dados. Elabore um breve resumo profissional, claro e objetivo, explicando:

            - Quais validações apresentaram divergências;
            - Quais foram validadas corretamente;
            - Destaque os pontos de atenção de forma sutil, sem alarmismo;
            - Evite repetir os mesmos termos;
            - Use parágrafos curtos e linguagem formal.

            Resultados:
        {texto_resultados}
    """

    modelo_IA = genai.GenerativeModel("gemini-1.5-flash-002")
    resposta = modelo_IA.generate_content(prompt)
    return resposta.text.strip()



def gerar_documento_word (resumo_validacao, arquivo_validacao_word):
    
    doc = Document(os.path.join(caminho_word_base, arquivo_word_base))

    data_atual = datetime.today().strftime("%d/%m/%Y")

    # CAPA
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '{{DATA}}' in cell.text:
                    cell.text = cell.text.replace('{{DATA}}', data_atual)
                
    # INTRODUÇÃO
    formata_titulo(doc, "1° INTRODUÇÃO", 2)
    for index, texto in introducao.items():
        formata_paragrafo(doc, texto)
    doc.add_paragraph("")

    # DADOS UTILIZADOS
    formata_titulo(doc, "2° DADOS UTILIZADOS", 2)
    formata_paragrafo(doc, "Arquivos utilizados para Apuração")

    for item in lista_dados_utilizados:
        doc.add_paragraph(f"• {item}")
    doc.add_paragraph("")

    # VALIDAÇÕES
    formata_titulo(doc, "3° VALIDAÇÕES", 2)
    doc.add_paragraph("")
    
    for i, (titulo, descricao) in enumerate(texto_validacoes.items(), start=1):
        doc.add_heading(f"3.{i} {titulo}", level = 3)
        formata_paragrafo(doc, descricao)

        qtde_pen = resumo_validacao.get(titulo, 0)  # Verifica a quantidade de divergência

        if qtde_pen > 0:
            resultado = f"Resultado: {status_text[1]}, com {qtde_pen} registros divergentes."
        else:
            resultado = f"Resultado: {status_text[0]}, sem divergências."

        formata_paragrafo(doc, resultado)
        doc.add_paragraph("")

    # RESUMO VALIDAÇÕES
    resumo_texto_ia = gerar_resumo_IA(resumo_validacao)

    formata_titulo(doc, "4° Resumo das Validações", 2)
    formata_paragrafo(doc, resumo_texto_ia)
    doc.add_paragraph("")

    doc.save(arquivo_validacao_word)


gerar_documento_word(status_item, arquivo_validacao_word)
        
    